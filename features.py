import re
import numpy as np
import pandas as pd
from docx import Document


class Vecteur:
    def __init__(self, vector) -> None:
        self.vec = vector
        self.compte = 0

    def compter(self):
        self.compte += 1


document = Document()


def document(texte: str):
    f = open(texte, "rb")
    doc = Document(f)
    f.close()
    return doc


# Création du vocabulaire
mot_interdit = ["le", "la", "l'", "de", "des", "un", "une", "d'", "à"]


def vocab(doc):
    vocabulaire = np.array([], dtype=object)
    phrases = []
    phrase = []
    paragraphe = doc.paragraphs
    for paragraph in paragraphe:
        for run in paragraph.runs:
            font = run.font
            couleur = font.color.rgb
            if couleur != None:
                couleur = (couleur[0], couleur[1], couleur[2])
            else:
                couleur = 0
            surlignage = font.highlight_color
            if not surlignage:
                surlignage = 0
            gras = font.bold
            if not gras:
                gras = 0
            italic = font.italic
            if not italic:
                italic = 0
            souligne = font.underline
            if not souligne:
                souligne = 0
            taille = font.size
            if not taille:
                taille = 0

            mots = re.split(r"\s+", run.text)
            for mot in mots:
                if mot not in mot_interdit:
                    prout = [mot, taille, gras, italic, souligne, couleur, surlignage]
                    if mot != "":
                        if prout not in [x.vec for x in vocabulaire]:
                            vecteur = Vecteur(prout)
                            vecteur.compter()
                            vocabulaire = np.append(vocabulaire, vecteur)
                        else:
                            kaka = vocabulaire[[x.vec == prout for x in vocabulaire]]
                            zizi = kaka[0]
                            zizi.compter()

                    phrase.append(prout)
                    if re.findall("[.|!|?]$", mot):
                        phrases.append(phrase)
                        phrase = []
    vocabulaire = sorted(vocabulaire, key=lambda x: x.compte)

    return vocabulaire, phrases


biero = document(
    "/Users/oscar/projet_tripro/Projet-Informatique-ML/Biero 20:06:23.docx"
)

vocabulaire, phrases = vocab(biero)

vocable = np.asanyarray([x.vec for x in vocabulaire], dtype=object)

print(vocable[:10])

import re
import numpy as np
import pandas as pd
from docx import Document
document = Document()

f = open("Test_2.docx", "rb")
doc = Document(f)
f.close()


def vecteur(doc):
    res=[]
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            #format de la run
            font=run.font
            couleur=font.color.rgb
            surligne=font.highlight_color
            taille=font.size
            gras=font.bold
            italic=font.italic
            souligne=font.underline
            police=font.name
            barre=font.strike
            caps=font.all_caps
            
            mots=re.split(r"\s+",run.text)
            for mot in mots:
                if mot != '':
                    res.append([mot,couleur,surligne,taille,gras,italic,souligne,police,barre,caps])
                
    return res
            


test = vecteur(doc)
test


def coef(vect):
    res=[]
    for x in vect:
        font=[]
        font.append(x[0])
        if x[1]==None:
            font.append((0,0,0))
        else:
            font.append((x[1][0],x[1][1],x[1][2]))
        for k in range(2,len(x)):
            if x[k]==None:
                font.append(0)
            elif x[k]==False:
                font.append(0)
            elif x[k]==True:
                font.append(1)
        res.append(font)
    return res


test2=coef(test)
test2


# ### Je vais faire une première distance pour tester
# #### Faudra normaliser pcq la composante couleur va bcp plus éloigner que les autres

# +
def dist_color(x,y):
    res=0
    for k in range(3):
        res+=(x[k]-y[k])**2
    return np.sqrt(res)

def distance(x,y):
    res=0
    res+=dist_color(x[1],x[2])
    for k in range(2,len(x)):
        res+=np.abs(x[k]-y[k])
    return res
    


# -

# ## Tableau pd tah grosse zoulette

test=np.array(test,dtype=object)
noms=test[:,0]
print(noms)
colonnes=['couleur','surligne','taille','gras','italic','souligne','police','barre','caps']

df = pd.DataFrame(test[:,1:],index=noms,columns=colonnes)

df

from docx import Document
import numpy as np
from features import vocable as vocabulaire, mot_interdit
import re

document = Document()


def document(texte: str):
    f = open(texte, "rb")
    doc = Document(f)
    f.close()
    return doc


doc = document("Test_2.docx")

new_doc = Document()


def lecture_vocab(one_hot: str):
    index = np.argmax(one_hot)
    return vocabulaire[index]


def mise_en_page(doc, mep: list, titre: str):
    k = 0
    for paragraph in doc.paragraphs:
        p = new_doc.add_paragraph()

        taille = 12
        gras = None
        italic = None
        souligne = None
        couleur = None
        surligne = None

        for run in paragraph.runs:
            mots = re.split(r"\s+", run.text)

            for mot in mots:
                r = p.add_run(mot)

                if mot not in mot_interdit:
                    vec = lecture_vocab(mep[k])
                    k += 1
                    taille = vec[1]
                    gras = vec[2]
                    italic = vec[3]
                    souligne = vec[4]
                    couleur = vec[5]
                    surligne = vec[6]

                r.font.size = taille
                r.bold = gras
                r.italic = italic
                r.font.underline = souligne
                r.font.color = couleur
                r.font.highlight_color = surligne

    new_doc.save(titre + ".docx")
